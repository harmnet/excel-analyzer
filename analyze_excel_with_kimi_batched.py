#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Excel数据集异常值分析程序

这个程序使用pandas读取Excel文件，然后将数据发送给Moonshot API，
让Kimi大模型分析数据集中的异常值，以表格形式返回异常字段的描述、问题分析和处理建议。
"""

import os
import sys
import time
import datetime
import json
import pandas as pd
import re
from pathlib import Path
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def show_spinner(seconds, message="处理中"):
    """显示加载动画"""
    spinner = "|/-\\"
    for _ in range(int(seconds * 5)):
        for char in spinner:
            sys.stdout.write(f"\r{message} {char}")
            sys.stdout.flush()
            time.sleep(0.2)
    sys.stdout.write("\r" + " " * (len(message) + 2) + "\r")
    sys.stdout.flush()

def num_tokens_from_string(string: str) -> int:
    """估算字符串中的token数量
    
    这是一个简单的估算方法，实际token数可能与OpenAI的分词器有差异
    中文和英文的token计算方式不同：
    - 英文通常是每4个字符约1个token
    - 中文通常是每1-2个字符约1个token
    """
    # 计算中文字符数
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', string))
    # 计算非中文字符数
    non_chinese_chars = len(string) - chinese_chars
    
    # 估算token数：中文字符按0.75个token/字符，非中文按0.25个token/字符
    estimated_tokens = chinese_chars * 0.75 + non_chinese_chars * 0.25
    
    return int(estimated_tokens)

def save_to_word(result_text, file_path, excel_file_name, model_name, process_time):
    """将分析结果保存为Word文档，并格式化表格内容"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('电子商务数据集异常分析结果', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加基本信息
    doc.add_paragraph(f"分析时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"分析文件: {excel_file_name}")
    doc.add_paragraph(f"使用模型: {model_name}")
    doc.add_paragraph(f"处理耗时: {process_time:.2f} 秒")
    
    doc.add_paragraph("=" * 50)
    
    # 处理文本内容，识别并格式化表格
    sections = result_text.split("### ")
    
    # 添加第一部分（如果有）
    if not result_text.startswith("### "):
        doc.add_paragraph(sections[0])
    
    # 处理其余部分
    for i in range(1 if result_text.startswith("### ") else 2, len(sections)):
        section = sections[i-1] if result_text.startswith("### ") else sections[i]
        
        # 添加标题
        doc.add_heading("### " + section.split("\n")[0], level=1)
        
        # 检查是否包含表格
        if "| 异常字段名称 | 异常值描述 | 问题分析 | 处理建议 |" in section:
            table_lines = []
            in_table = False
            
            for line in section.split("\n"):
                if line.startswith("|"):
                    in_table = True
                    table_lines.append(line)
                elif in_table and line.strip() == "":
                    in_table = False
                    
                    # 处理表格
                    if len(table_lines) >= 3:  # 表头 + 分隔行 + 至少一行数据
                        # 获取列数
                        columns = len(table_lines[0].split("|")) - 2  # 减去开头和结尾的|
                        
                        # 创建表格
                        table = doc.add_table(rows=len(table_lines)-1, cols=columns)
                        table.style = 'Table Grid'
                        
                        # 填充表头
                        header_cells = table_lines[0].split("|")[1:-1]
                        for j, cell_text in enumerate(header_cells):
                            table.cell(0, j).text = cell_text.strip()
                        
                        # 填充数据行
                        for row_idx, line in enumerate(table_lines[2:], start=1):
                            cells = line.split("|")[1:-1]
                            for col_idx, cell_text in enumerate(cells):
                                table.cell(row_idx, col_idx).text = cell_text.strip()
                    
                    table_lines = []
                elif not in_table:
                    # 添加普通段落
                    if line.strip() and not line.startswith("### "):
                        doc.add_paragraph(line)
            
            # 如果表格在文本末尾，确保它也被处理
            if in_table and len(table_lines) >= 3:
                # 获取列数
                columns = len(table_lines[0].split("|")) - 2
                
                # 创建表格
                table = doc.add_table(rows=len(table_lines)-1, cols=columns)
                table.style = 'Table Grid'
                
                # 填充表头
                header_cells = table_lines[0].split("|")[1:-1]
                for j, cell_text in enumerate(header_cells):
                    table.cell(0, j).text = cell_text.strip()
                
                # 填充数据行
                for row_idx, line in enumerate(table_lines[2:], start=1):
                    cells = line.split("|")[1:-1]
                    for col_idx, cell_text in enumerate(cells):
                        table.cell(row_idx, col_idx).text = cell_text.strip()
        else:
            # 没有表格，直接添加文本内容
            content_lines = section.split("\n")[1:]  # 跳过标题行
            for line in content_lines:
                if line.strip():
                    doc.add_paragraph(line)
    
    # 保存文档
    doc.save(file_path)
    return file_path

def main():
    print("="*50)
    print("Excel数据集异常值分析程序 (分批处理版)")
    print("="*50)
    print("本程序将使用Kimi大模型分析Excel数据集中的异常值")
    print("数据将分批发送给模型，每批后等待30秒，避免超出API限制")
    print("程序会显示每批数据的文字数量和估算的Tokens数量")
    print()
    
    try:
        # 检查是否安装了python-docx
        try:
            import docx
        except ImportError:
            print("正在安装python-docx库，用于生成Word文档...")
            import subprocess
            subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
            print("安装完成！")
        
        # 设置API密钥
        api_key = "sk-uMgSk5k16xVya8DMUSpBH0inGDg0v42XT6bKLuqlTkbcJtTA"
        
        # 获取用户主目录
        home_dir = Path(os.path.expanduser("~"))
        
        # 构建Excel文件的完整路径
        excel_file_path = home_dir / "Desktop" / "《商务数据分析基础》数据集-电子商务-完整数据.xlsx"
        
        # 检查文件是否存在
        if not excel_file_path.exists():
            print(f"错误：文件 {excel_file_path} 不存在！")
            print("请确认文件名称是否正确，以及文件是否位于桌面上")
            return
        
        print(f"找到文件：{excel_file_path}")
        print("文件大小：{:.2f} MB".format(excel_file_path.stat().st_size / (1024 * 1024)))
        
        # 使用pandas读取Excel文件
        print("\n正在读取Excel文件...")
        show_spinner(1, "读取中")
        
        try:
            # 读取Excel文件
            df = pd.read_excel(excel_file_path)
            print(f"成功读取Excel文件，共有 {len(df)} 行，{len(df.columns)} 列")
            
            # 显示数据预览
            print("\n数据预览 (前5行):")
            print(df.head().to_string())
            
            # 显示数据类型
            print("\n数据类型:")
            print(df.dtypes)
            
            # 将数据转换为文本格式
            print("\n正在准备数据...")
            show_spinner(1, "准备中")
            
            # 准备数据集基本信息
            data_info = "数据集信息：\n"
            data_info += f"- 行数：{len(df)}\n"
            data_info += f"- 列数：{len(df.columns)}\n"
            data_info += f"- 列名：{', '.join(df.columns.tolist())}\n\n"
            
            # 添加数据类型信息
            data_info += "数据类型：\n"
            for col, dtype in df.dtypes.items():
                data_info += f"- {col}: {dtype}\n"
            
            # 分批处理数据
            batch_size = 600  # 每批600行
            num_batches = (len(df) + batch_size - 1) // batch_size
            
            print(f"\n数据将被分为 {num_batches} 批处理，每批最多 {batch_size} 行")
            print("每批数据发送后将等待30秒，避免超出API限制")
            
            print("\n正在初始化Moonshot API客户端...")
            show_spinner(1, "初始化中")
            
            # 初始化OpenAI客户端
            client = OpenAI(
                api_key=api_key,
                base_url="https://api.moonshot.cn/v1"
            )
            
            # 设置使用的模型
            model_name = "moonshot-v1-128k"
            print(f"\n使用模型：{model_name}")
            print("该模型拥有128K的上下文窗口，可以处理更大的文件和更复杂的分析任务")
            
            # 构建初始系统消息
            system_message = """你是 Kimi，由 Moonshot AI 提供的人工智能助手。你将分析一个Excel数据集，寻找其中的异常值。

重要说明：由于数据量较大，我将分多个批次发送数据给你。每个批次发送后，你只需回复"已接收数据批次X/Y"，不需要进行分析。
当所有批次都发送完毕后，我会发送最终的分析请求，那时你再进行完整的分析。"""
            
            # 准备消息历史
            conversation_history = [
                {"role": "system", "content": system_message},
                {"role": "system", "content": data_info}
            ]
            
            # 分批发送数据并等待
            for i in range(num_batches):
                start_idx = i * batch_size
                end_idx = min((i + 1) * batch_size, len(df))
                
                batch_df = df.iloc[start_idx:end_idx]
                
                # 转换为文本格式
                batch_text = f"\n数据批次 {i+1}/{num_batches} (行 {start_idx+1} 到 {end_idx})：\n"
                batch_text += batch_df.to_string(max_rows=None, max_cols=None)
                
                # 计算文字数量和估算的tokens数量
                chars_count = len(batch_text)
                tokens_estimate = num_tokens_from_string(batch_text)
                
                print(f"\n准备发送数据批次 {i+1}/{num_batches} (行 {start_idx+1} 到 {end_idx})")
                print(f"批次大小: {chars_count} 字符, 估算 {tokens_estimate} tokens")
                
                # 发送当前批次数据
                try:
                    print(f"正在发送批次 {i+1}/{num_batches}...")
                    
                    # 构建当前批次的消息
                    current_messages = conversation_history.copy()
                    current_messages.append({"role": "user", "content": f"这是数据批次 {i+1}/{num_batches}:\n{batch_text}"})
                    
                    # 计算当前请求的总tokens估算
                    total_chars = sum(len(msg["content"]) for msg in current_messages)
                    total_tokens = num_tokens_from_string(json.dumps([msg["content"] for msg in current_messages], ensure_ascii=False))
                    print(f"当前请求总大小: {total_chars} 字符, 估算 {total_tokens} tokens")
                    
                    # 发送请求
                    response = client.chat.completions.create(
                        model=model_name,
                        messages=current_messages,
                        temperature=0.2
                    )
                    
                    # 获取回复
                    batch_response = response.choices[0].message.content
                    print(f"模型回复: {batch_response}")
                    
                    # 如果API返回了token使用信息，显示出来
                    if hasattr(response, 'usage') and response.usage:
                        print(f"API报告的token使用情况:")
                        print(f"- 输入tokens: {response.usage.prompt_tokens}")
                        print(f"- 输出tokens: {response.usage.completion_tokens}")
                        print(f"- 总tokens: {response.usage.total_tokens}")
                    
                    # 更新对话历史
                    conversation_history.append({"role": "user", "content": f"这是数据批次 {i+1}/{num_batches}:\n{batch_text}"})
                    conversation_history.append({"role": "assistant", "content": batch_response})
                    
                    # 如果不是最后一批，等待30秒再继续
                    if i < num_batches - 1:
                        print(f"等待30秒后发送下一批数据...")
                        for j in range(30, 0, -1):
                            print(f"\r等待中... {j} 秒", end="", flush=True)
                            time.sleep(1)
                        print("\r等待完成，继续发送...                 ")
                    
                except Exception as e:
                    error_str = str(e)
                    
                    # 打印详细的错误信息
                    print(f"\n发送批次 {i+1} 时出错:")
                    print(f"错误详情: {error_str}")
                    print(f"错误类型: {type(e).__name__}")
                    
                    # 如果是API错误，尝试提取更多信息
                    if hasattr(e, 'response'):
                        try:
                            error_response = e.response
                            print(f"API错误状态码: {error_response.status_code}")
                            error_json = error_response.json()
                            print(f"API错误信息: {json.dumps(error_json, ensure_ascii=False, indent=2)}")
                        except:
                            print("无法解析API错误响应")
                    
                    # 检查是否是频率限制错误
                    if "rate_limit_reached_error" in error_str:
                        print(f"调用API时遇到频率限制，等待45秒后重试...")
                        
                        # 倒计时显示
                        for j in range(45, 0, -1):
                            print(f"\r等待中... {j} 秒", end="", flush=True)
                            time.sleep(1)
                        print("\r等待完成，重试发送...                 ")
                        
                        # 重试当前批次
                        i -= 1  # 回退一步，重试当前批次
                        continue
                    
                    # 其他错误，终止程序
                    print("发送数据时出错，程序终止")
                    return
            
            # 所有批次发送完毕，发送最终分析请求
            print("\n所有数据批次已发送完毕，现在请求模型进行分析...")
            
            # 构建查询提示
            query = """
我已经将完整的数据集分批发送给你。现在请详细分析这个Excel数据集中哪些字段的值存在异常。

请以表格形式提供以下信息：
1. 异常字段名称
2. 异常值描述（包括异常值的范围、类型或具体例子）
3. 问题分析（为什么这些值被认为是异常的，可能的原因）
4. 处理建议（如何清洗或修正这些异常值的具体方法）

表格应该包含这四列，并为每个发现的异常字段提供一行数据。
如果没有发现异常，请说明数据集看起来正常。

在回答前，请先简要描述一下这个数据集的基本情况，包括：
- 数据集的行数和列数
- 各列的数据类型
- 数据集的主要内容和用途
- 数据的时间范围（如果适用）

最后，请提供一些关于这个数据集整体质量的建议，以及如何更好地利用这些数据进行商务分析。
"""
            
            # 添加最终的用户查询到对话历史
            conversation_history.append({"role": "user", "content": query})
            
            # 计算最终请求的总tokens估算
            final_request_chars = sum(len(msg["content"]) for msg in conversation_history)
            final_request_tokens = num_tokens_from_string(json.dumps([msg["content"] for msg in conversation_history], ensure_ascii=False))
            print(f"最终分析请求总大小: {final_request_chars} 字符, 估算 {final_request_tokens} tokens")
            
            # 发送聊天请求
            print("\n正在请求Kimi分析数据集...")
            print(f"(使用{model_name}模型分析，过程可能需要几分钟时间，请耐心等待)")
            
            start_time = time.time()
            
            # 添加重试机制
            max_retries = 3
            retry_count = 0
            retry_delay = 45
            
            while retry_count < max_retries:
                try:
                    print(f"正在发送分析请求... (尝试 {retry_count + 1}/{max_retries})")
                    
                    # 使用完整的对话历史发送请求
                    response = client.chat.completions.create(
                        model=model_name,
                        messages=conversation_history,
                        temperature=0.2
                    )
                    
                    # 计算处理时间
                    process_time = time.time() - start_time
                    
                    # 获取Kimi的回答内容
                    analysis_result = response.choices[0].message.content
                    
                    # 如果API返回了token使用信息，显示出来
                    if hasattr(response, 'usage') and response.usage:
                        print(f"API报告的最终分析token使用情况:")
                        print(f"- 输入tokens: {response.usage.prompt_tokens}")
                        print(f"- 输出tokens: {response.usage.completion_tokens}")
                        print(f"- 总tokens: {response.usage.total_tokens}")
                    
                    # 打印Kimi的回答
                    print("\n" + "="*80 + "\n")
                    print("Kimi的数据集异常分析结果：\n")
                    print(analysis_result)
                    print("\n" + "="*80)
                    print(f"\n分析耗时：{process_time:.2f} 秒")
                    
                    # 保存分析结果到Word文件
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    result_file_name = f"电子商务数据集异常分析_{model_name}_{timestamp}.docx"
                    result_file_path = home_dir / "Desktop" / result_file_name
                    
                    # 保存为Word文档
                    saved_path = save_to_word(
                        analysis_result, 
                        result_file_path, 
                        excel_file_path.name, 
                        model_name, 
                        process_time
                    )
                    
                    print(f"\n分析结果已保存到: {saved_path}")
                    
                    # 成功获取结果，跳出循环
                    break
                    
                except Exception as e:
                    error_str = str(e)
                    retry_count += 1
                    
                    # 打印详细的错误信息
                    print(f"\n错误详情: {error_str}")
                    print(f"错误类型: {type(e).__name__}")
                    
                    # 如果是API错误，尝试提取更多信息
                    if hasattr(e, 'response'):
                        try:
                            error_response = e.response
                            print(f"API错误状态码: {error_response.status_code}")
                            error_json = error_response.json()
                            print(f"API错误信息: {json.dumps(error_json, ensure_ascii=False, indent=2)}")
                        except:
                            print("无法解析API错误响应")
                    
                    # 检查是否是频率限制错误
                    if "rate_limit_reached_error" in error_str:
                        # 固定等待45秒
                        wait_time = 45
                        
                        print(f"调用API时遇到频率限制，等待 {wait_time} 秒后重试... (尝试 {retry_count}/{max_retries})")
                        
                        # 倒计时显示
                        for i in range(wait_time, 0, -1):
                            print(f"\r等待中... {i} 秒", end="", flush=True)
                            time.sleep(1)
                        print("\r等待完成，继续请求...                 ")
                        continue
                    
                    # 如果不是频率限制错误，或者已经达到最大重试次数，则抛出异常
                    if retry_count >= max_retries:
                        raise Exception(f"调用API失败: 达到最大重试次数 ({max_retries})")
                    
                    print(f"调用API失败，45秒后重试: {error_str}")
                    time.sleep(45)
            
            # 如果所有重试都失败，抛出异常
            if retry_count >= max_retries:
                raise Exception("调用API失败: 达到最大重试次数")
                
        except pd.errors.EmptyDataError:
            print("错误：Excel文件为空")
        except pd.errors.ParserError:
            print("错误：无法解析Excel文件，文件可能已损坏")
        except Exception as e:
            print(f"分析过程中出错：{str(e)}")
            print("请检查您的API密钥是否正确，以及网络连接是否正常")
        
    except KeyboardInterrupt:
        print("\n\n程序被用户中断")
    except Exception as e:
        print(f"\n程序运行出错：{str(e)}")
    finally:
        print("\n程序执行完毕")

if __name__ == "__main__":
    main()