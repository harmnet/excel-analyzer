#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Excel数据集异常值分析GUI程序

这个程序提供一个图形用户界面，让用户可以：
1. 选择本地Excel文件
2. 设置分析参数（批次大小、等待时间等）
3. 一键启动AI分析
4. 查看分析进度和结果
5. 将分析结果保存为Word文档
"""

import os
import sys
import time
import datetime
import json
import pandas as pd
import re
import threading
from pathlib import Path
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# PyQt5导入
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, 
    QPushButton, QLabel, QFileDialog, QTextEdit, QProgressBar, 
    QSpinBox, QGroupBox, QFormLayout, QMessageBox, QTabWidget,
    QLineEdit, QCheckBox, QComboBox, QSplitter, QFrame, QDialog,
    QScrollArea, QStyleFactory, QToolButton, QStyle
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QSize, QObject
from PyQt5.QtGui import QFont, QIcon, QTextCursor, QColor, QPalette, QBrush, QPixmap

# 常量定义
DEFAULT_MOONSHOT_API_KEY = "sk-uMgSk5k16xVya8DMUSpBH0inGDg0v42XT6bKLuqlTkbcJtTA"
DEFAULT_SILICONFLOW_API_KEY = "sk-bivnwauskdbvpspvmdorrgkrpwlyfxbfcezqsfsevowzubdj"  # 请填入您的硅基流动API密钥(纯密钥格式如"sk-xxxx"，不需要添加Bearer前缀)
DEFAULT_VOLCANO_API_KEY = "5bcbac4d-6ce8-4f41-b27d-63764f7c74f6"  # 火山引擎API密钥
DEFAULT_BATCH_SIZE = 600
DEFAULT_WAIT_TIME = 10
DEFAULT_MODEL = "Pro/deepseek-ai/DeepSeek-R1"  # 使用官方文档支持的模型名称

# 颜色定义
COLOR_BURGUNDY = QColor(158, 33, 65)  # 紫荆红
COLOR_GOLD = QColor(177, 136, 101)    # 金色
COLOR_GRAY = QColor(166, 166, 166)    # 灰色

# 模型配置
MODEL_CONFIG = {
    "moonshot-v1-32k": {
        "token_limit": 32768,
        "api_type": "moonshot",
        "api_key": DEFAULT_MOONSHOT_API_KEY,
        "batch_size_suggestion": 600,
        "description": "Moonshot AI 32K上下文",
        "icon": "🌙"
    },
    "moonshot-v1-128k": {
        "token_limit": 131072,
        "api_type": "moonshot", 
        "api_key": DEFAULT_MOONSHOT_API_KEY,
        "batch_size_suggestion": 2000,
        "description": "Moonshot AI 128K上下文",
        "icon": "🌟"
    },
    "Pro/deepseek-ai/DeepSeek-R1": {
        "token_limit": 65536,
        "api_type": "siliconflow",
        "api_key": DEFAULT_SILICONFLOW_API_KEY,
        "batch_size_suggestion": 1200,
        "description": "DeepSeek-R1 64K上下文",
        "icon": "🔍"
    },
    "Pro/deepseek-ai/DeepSeek-V3": {
        "token_limit": 65536,
        "api_type": "siliconflow",
        "api_key": DEFAULT_SILICONFLOW_API_KEY,
        "batch_size_suggestion": 1200,
        "description": "DeepSeek-V3 64K上下文",
        "icon": "🔍"
    },
    "Qwen/QwQ-32B": {
        "token_limit": 32768,
        "api_type": "siliconflow",
        "api_key": DEFAULT_SILICONFLOW_API_KEY,
        "batch_size_suggestion": 100,  # 更保守的批次大小设置
        "description": "千问QwQ-32B 32K上下文",
        "icon": "🔍"
    },
    "deepseek-r1-250120": {
        "token_limit": 65536,
        "api_type": "volcano",
        "api_key": DEFAULT_VOLCANO_API_KEY,
        "batch_size_suggestion": 1200,
        "description": "火山引擎 DeepSeek-R1 64K上下文",
        "icon": "🌋"
    }
}

# 如果硅基流动API模块存在，导入它
try:
    from silicon_flow_api import SiliconFlowAPI
    SILICON_FLOW_AVAILABLE = True
except ImportError:
    SILICON_FLOW_AVAILABLE = False

# 添加一个简化版的SiliconFlowAPI类，专门用于处理400错误情况
class SimplifiedSiliconFlowAPI:
    """简化版的硅基流动API调用类，专门用于处理400错误情况"""
    
    def __init__(self, api_key, base_url="https://api.siliconflow.cn/v1"):
        """初始化API客户端"""
        self.api_key = api_key
        self.base_url = base_url
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
    
    def chat_completion(self, messages, model="Pro/deepseek-ai/DeepSeek-R1", temperature=0.7, max_tokens=2000):
        """简化版的聊天接口，只使用必要的参数"""
        import requests
        import json
        
        url = f"{self.base_url}/chat/completions"
        
        # 获取模型的token限制
        model_token_limit = 65536  # DeepSeek模型默认65536
        if "32k" in model.lower() or "qwen" in model.lower():
            model_token_limit = 32768
        elif "128k" in model.lower():
            model_token_limit = 131072
        
        # 确保max_tokens不会导致总token超过限制
        if "qwen" in model.lower():
            # 对于Qwen模型，更加保守地限制输出token
            safe_max_tokens = min(max_tokens, 2000)  # 限制输出token不超过2000
        else:
            safe_max_tokens = min(max_tokens, 4000)  # 限制输出token不超过4000
        
        # 极度简化参数，只使用最基本必需的参数
        payload = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": safe_max_tokens
        }
        
        # 打印请求详情用于调试
        print(f"\n发送请求到: {url}")
        print(f"使用模型: {model}")
        print(f"API密钥格式: {self.api_key[:5]}...{self.api_key[-4:]}")
        print(f"消息数量: {len(messages)}")
        print(f"最大输出tokens: {safe_max_tokens}")
        print(f"请求体: {json.dumps(payload, ensure_ascii=False)[:200]}...")
        
        try:
            response = requests.post(url, json=payload, headers=self.headers)
            
            # 如果请求失败，打印详细信息
            if not response.ok:
                print(f"API请求失败，状态码: {response.status_code}")
                print(f"响应内容: {response.text}")
                
                # 尝试解析JSON响应
                try:
                    error_json = response.json()
                    print(f"错误详情: {json.dumps(error_json, ensure_ascii=False, indent=2)}")
                except:
                    print("响应内容不是有效的JSON")
            
            response.raise_for_status()  # 抛出HTTP错误
            return response.json()
        except requests.exceptions.RequestException as e:
            raise Exception(f"API请求失败: {str(e)}")

class WorkerSignals(QObject):
    """
    定义工作线程的信号
    """
    update_progress = pyqtSignal(str)
    update_progress_bar = pyqtSignal(int)
    analysis_complete = pyqtSignal(str, str)
    error = pyqtSignal(str)
    
class AnalysisWorker(QThread):
    """
    执行Excel分析的工作线程
    """
    def __init__(self, file_path, api_key, batch_size, wait_time, model_name, max_rows=None):
        super().__init__()
        self.file_path = file_path
        self.api_key = api_key
        self.batch_size = batch_size
        self.wait_time = wait_time
        self.model_name = model_name
        self.max_rows = max_rows
        self.signals = WorkerSignals()
        self.stop_requested = False
        
    def num_tokens_from_string(self, string):
        """估算字符串中的token数量"""
        # 计算中文字符数
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', string))
        # 计算非中文字符数
        non_chinese_chars = len(string) - chinese_chars
        
        # 估算token数：中文字符按0.75个token/字符，非中文按0.25个token/字符
        estimated_tokens = chinese_chars * 0.75 + non_chinese_chars * 0.25
        
        # 为Qwen模型应用更高的系数，补偿低估现象
        if hasattr(self, 'model_name') and 'qwen' in self.model_name.lower():
            # 为Qwen模型应用2.2倍的系数，确保不会低估
            estimated_tokens *= 2.2
            
        return int(estimated_tokens)
        
    def save_to_word(self, result_text, excel_file_name, process_time):
        """将分析结果保存为Word文档，并格式化表格内容"""
        doc = Document()
        
        # 添加标题
        title = doc.add_heading('Excel数据集异常分析结果', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息
        doc.add_paragraph(f"分析时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"分析文件: {excel_file_name}")
        doc.add_paragraph(f"使用模型: {self.model_name}")
        doc.add_paragraph(f"处理耗时: {process_time:.2f} 秒")
        
        doc.add_paragraph("=" * 50)
        
        # 处理文本内容，识别并格式化表格
        sections = result_text.split("### ")
        
        # 添加第一部分（如果有）
        if not result_text.startswith("### "):
            doc.add_paragraph(sections[0])
        
        # 处理其余部分
        for i in range(1 if result_text.startswith("### ") else 2, len(sections)):
            section = sections[i-1] if result_text.startswith("### ") else sections[i]
            
            # 添加标题
            doc.add_heading("### " + section.split("\n")[0], level=1)
            
            # 检查是否包含表格
            if "| 异常字段名称 | 异常值描述 | 问题分析 | 处理建议 |" in section:
                table_lines = []
                in_table = False
                
                for line in section.split("\n"):
                    if line.startswith("|"):
                        in_table = True
                        table_lines.append(line)
                    elif in_table and line.strip() == "":
                        in_table = False
                        
                        # 处理表格
                        if len(table_lines) >= 3:  # 表头 + 分隔行 + 至少一行数据
                            # 获取列数
                            columns = len(table_lines[0].split("|")) - 2  # 减去开头和结尾的|
                            
                            # 创建表格
                            table = doc.add_table(rows=len(table_lines)-1, cols=columns)
                            table.style = 'Table Grid'
                            
                            # 填充表头
                            header_cells = table_lines[0].split("|")[1:-1]
                            for j, cell_text in enumerate(header_cells):
                                table.cell(0, j).text = cell_text.strip()
                            
                            # 填充数据行
                            for row_idx, line in enumerate(table_lines[2:], start=1):
                                cells = line.split("|")[1:-1]
                                for col_idx, cell_text in enumerate(cells):
                                    table.cell(row_idx, col_idx).text = cell_text.strip()
                        
                        table_lines = []
                    elif not in_table:
                        # 添加普通段落
                        if line.strip() and not line.startswith("### "):
                            doc.add_paragraph(line)
                
                # 如果表格在文本末尾，确保它也被处理
                if in_table and len(table_lines) >= 3:
                    # 获取列数
                    columns = len(table_lines[0].split("|")) - 2
                    
                    # 创建表格
                    table = doc.add_table(rows=len(table_lines)-1, cols=columns)
                    table.style = 'Table Grid'
                    
                    # 填充表头
                    header_cells = table_lines[0].split("|")[1:-1]
                    for j, cell_text in enumerate(header_cells):
                        table.cell(0, j).text = cell_text.strip()
                    
                    # 填充数据行
                    for row_idx, line in enumerate(table_lines[2:], start=1):
                        cells = line.split("|")[1:-1]
                        for col_idx, cell_text in enumerate(cells):
                            table.cell(row_idx, col_idx).text = cell_text.strip()
            else:
                # 没有表格，直接添加文本内容
                content_lines = section.split("\n")[1:]  # 跳过标题行
                for line in content_lines:
                    if line.strip():
                        doc.add_paragraph(line)
        
        # 确定保存路径
        home_dir = Path(os.path.expanduser("~"))
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # 处理模型名中的斜杠，将其替换为下划线，避免文件名中出现非法字符
        safe_model_name = self.model_name.replace("/", "_").replace("\\", "_")
        
        result_file_name = f"Excel数据集异常分析_{safe_model_name}_{timestamp}.docx"
        desktop_dir = home_dir / "Desktop"
        
        # 确保目标目录存在
        if not desktop_dir.exists():
            self.signals.update_progress.emit(f"桌面目录不存在，将保存到用户主目录: {home_dir}")
            result_file_path = home_dir / result_file_name
        else:
            result_file_path = desktop_dir / result_file_name
        
        # 确保保存成功
        try:
            doc.save(str(result_file_path))  # 使用str确保兼容性
            self.signals.update_progress.emit(f"文件已成功保存到: {result_file_path}")
        except Exception as e:
            self.signals.update_progress.emit(f"保存到 {result_file_path} 失败: {str(e)}")
            # 尝试保存到用户主目录
            fallback_path = home_dir / result_file_name
            self.signals.update_progress.emit(f"尝试保存到用户主目录: {fallback_path}")
            doc.save(str(fallback_path))
            result_file_path = fallback_path
        
        return result_file_path
    
    def direct_api_request(self, messages, api_key, model_name="Pro/deepseek-ai/DeepSeek-R1"):
        """直接发送API请求的备用方法，不使用任何API库，直接使用requests"""
        import requests
        import json
        import time
        
        # 使用官方文档中的样例URL和最简化的请求体
        url = "https://api.siliconflow.cn/v1/chat/completions"
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # 从官方文档示例中提取的最小化请求参数
        payload = {
            "model": model_name,
            "messages": messages
        }
        
        # 记录请求信息
        self.signals.update_progress.emit(f"\n--- 直接API请求方法 ---")
        self.signals.update_progress.emit(f"请求URL: {url}")
        self.signals.update_progress.emit(f"使用模型: {model_name}")
        self.signals.update_progress.emit(f"API密钥格式: {api_key[:5]}...{api_key[-4:]}")
        self.signals.update_progress.emit(f"消息数量: {len(messages)}")
        
        # 发送请求
        try:
            self.signals.update_progress.emit("发送请求...")
            start_time = time.time()
            response = requests.post(url, json=payload, headers=headers)
            elapsed_time = time.time() - start_time
            
            self.signals.update_progress.emit(f"响应时间: {elapsed_time:.2f}秒")
            self.signals.update_progress.emit(f"状态码: {response.status_code}")
            
            # 检查响应
            if response.status_code == 200:
                self.signals.update_progress.emit("请求成功！")
                response_json = response.json()
                
                # 提取消息内容
                if "choices" in response_json and len(response_json["choices"]) > 0:
                    message = response_json["choices"][0].get("message", {})
                    content = message.get("content", "")
                    self.signals.update_progress.emit(f"模型回复长度: {len(content)} 字符")
                    
                    # 记录token使用情况
                    if "usage" in response_json:
                        usage = response_json["usage"]
                        self.signals.update_progress.emit(f"Token使用情况:")
                        self.signals.update_progress.emit(f"- 输入tokens: {usage.get('prompt_tokens', 'N/A')}")
                        self.signals.update_progress.emit(f"- 输出tokens: {usage.get('completion_tokens', 'N/A')}")
                        self.signals.update_progress.emit(f"- 总tokens: {usage.get('total_tokens', 'N/A')}")
                    
                    return response_json
                else:
                    self.signals.update_progress.emit("响应缺少choices字段")
                    self.signals.update_progress.emit(f"完整响应: {json.dumps(response_json, ensure_ascii=False, indent=2)[:500]}...")
                    raise Exception("API响应格式错误")
            else:
                # 尝试解析错误响应
                self.signals.update_progress.emit(f"请求失败: HTTP {response.status_code}")
                try:
                    error_json = response.json()
                    error_str = json.dumps(error_json, ensure_ascii=False, indent=2)
                    self.signals.update_progress.emit(f"错误详情: {error_str}")
                    
                    # 检查是否与模型相关的错误
                    error_message = str(error_json.get("error", {}).get("message", "")).lower()
                    if ("model" in error_message or 
                        "not found" in error_message or 
                        "does not exist" in error_message or
                        "invalid" in error_message):
                        
                        # 首先尝试DeepSeek-V3，如果当前模型不是DeepSeek-V3且不是QwQ-32B
                        if model_name != "Pro/deepseek-ai/DeepSeek-V3" and model_name != "Qwen/QwQ-32B":
                            self.signals.update_progress.emit(f"⚠️ 模型错误: {model_name} 似乎不可用")
                            self.signals.update_progress.emit(f"尝试使用备选模型 Pro/deepseek-ai/DeepSeek-V3...")
                            return self.direct_api_request(messages, api_key, "Pro/deepseek-ai/DeepSeek-V3")
                        
                        # 如果当前是DeepSeek-V3但失败了，尝试QwQ-32B
                        elif model_name == "Pro/deepseek-ai/DeepSeek-V3" and model_name != "Qwen/QwQ-32B":
                            self.signals.update_progress.emit(f"⚠️ 模型错误: {model_name} 似乎不可用")
                            self.signals.update_progress.emit(f"尝试使用默认模型 Qwen/QwQ-32B...")
                            return self.direct_api_request(messages, api_key, "Qwen/QwQ-32B")
                except:
                    self.signals.update_progress.emit(f"原始响应: {response.text[:500]}")
                
                raise Exception(f"API请求失败: HTTP {response.status_code}")
        
        except Exception as e:
            self.signals.update_progress.emit(f"请求异常: {str(e)}")
            raise
    
    def run(self):
        """执行分析任务"""
        try:
            # 更新进度信息
            self.signals.update_progress.emit("正在开始分析任务...")
            self.signals.update_progress_bar.emit(5)
            
            # 检查文件是否存在
            file_path = Path(self.file_path)
            if not file_path.exists():
                self.signals.error.emit(f"错误：文件 {file_path} 不存在！")
                return
            
            self.signals.update_progress.emit(f"找到文件：{file_path}")
            self.signals.update_progress.emit(f"文件大小：{file_path.stat().st_size / (1024 * 1024):.2f} MB")
            self.signals.update_progress_bar.emit(10)
            
            # 使用pandas读取Excel文件
            self.signals.update_progress.emit("\n正在读取Excel文件...")
            
            try:
                # 读取Excel文件
                df = pd.read_excel(file_path)
                
                # 如果指定了最大行数，则截取
                if self.max_rows and len(df) > self.max_rows:
                    self.signals.update_progress.emit(f"由于行数较多，将只分析前 {self.max_rows} 行数据")
                    df = df.iloc[:self.max_rows]
                    
                self.signals.update_progress.emit(f"成功读取Excel文件，共有 {len(df)} 行，{len(df.columns)} 列")
                
                # 注释掉数据预览显示
                # self.signals.update_progress.emit("\n数据预览 (前5行):")
                # self.signals.update_progress.emit(df.head().to_string())
                
                # 注释掉数据类型显示
                # self.signals.update_progress.emit("\n数据类型:")
                # self.signals.update_progress.emit(str(df.dtypes))
                
                self.signals.update_progress_bar.emit(20)
                
                # 准备数据集基本信息
                data_info = "数据集信息：\n"
                data_info += f"- 行数：{len(df)}\n"
                data_info += f"- 列数：{len(df.columns)}\n"
                data_info += f"- 列名：{', '.join(df.columns.tolist())}\n\n"
                
                # 添加数据类型信息
                data_info += "数据类型：\n"
                for col, dtype in df.dtypes.items():
                    data_info += f"- {col}: {dtype}\n"
                
                # 获取模型配置
                model_info = MODEL_CONFIG.get(self.model_name, MODEL_CONFIG[DEFAULT_MODEL])
                model_token_limit = model_info.get("token_limit", 32768)
                model_api_type = model_info.get("api_type", "moonshot")
                
                # 先构建初始系统消息
                system_message = """你是一位数据分析助手。你将分析一个Excel数据集，寻找其中的异常值。

重要说明：由于数据量较大，数据将分多个批次发送给你。每个批次发送后，你只需回复"已接收数据批次X/Y"，不需要进行分析。
当所有批次都发送完毕后，将发送最终的分析请求，那时你再进行完整的分析。"""
                
                # 准备消息历史
                initial_system_messages = [
                    {"role": "system", "content": system_message},
                    {"role": "system", "content": data_info}
                ]
                
                # 动态调整批次大小
                batch_size = self.batch_size
                if batch_size <= 0:  # 如果批次大小未设置，自动计算最佳批次大小
                    # 预估每行数据平均token数
                    sample_size = min(50, len(df))  # 取前50行或所有行作为样本
                    sample_df = df.iloc[:sample_size]
                    sample_text = sample_df.to_string(max_rows=None, max_cols=None)
                    sample_tokens = self.num_tokens_from_string(sample_text)
                    avg_tokens_per_row = sample_tokens / sample_size
                    
                    # 计算系统消息和提示的token数
                    system_tokens = sum(self.num_tokens_from_string(msg["content"]) for msg in initial_system_messages)
                    
                    # 更保守地计算可用tokens，确保不超过模型的真实限制
                    # 硅基流动API要求总token数必须小于模型的max_seq_len
                    if "qwen" in self.model_name.lower():
                        # Qwen/QwQ-32B需要更保守的处理，将安全限制降低到24000
                        safe_token_limit = min(24000, model_token_limit - 3000)  # 为Qwen模型留出3000 tokens的安全边界
                        # self.signals.update_progress.emit(f"注意：检测到使用的是Qwen模型，将使用更保守的token限制")
                    else:
                        safe_token_limit = min(64000, model_token_limit - 1000)  # 留出1000 tokens的安全边界
                    
                    # 为API响应和其他开销预留空间（约25%）
                    available_tokens = int(safe_token_limit * 0.75) - system_tokens
                    
                    # 计算每批次可以处理的最大行数
                    max_rows_per_batch = int(available_tokens / avg_tokens_per_row)
                    
                    # 额外减少行数作为安全余量
                    if "qwen" in self.model_name.lower():
                        # Qwen模型额外减少30%的行数作为安全余量
                        max_rows_per_batch = int(max_rows_per_batch * 0.7)
                        # self.signals.update_progress.emit(f"- 为Qwen模型额外应用了30%的安全系数")
                    else:
                        # 其他模型减少5%的行数作为安全余量
                        max_rows_per_batch = int(max_rows_per_batch * 0.95)
                    
                    # 确保每批至少处理10行，避免批次过小
                    max_rows_per_batch = max(10, max_rows_per_batch)
                    
                    # 使用计算出的批次大小
                    batch_size = max_rows_per_batch
                    
                    # 注释掉自动计算批次大小相关的日志输出
                    # self.signals.update_progress.emit(f"自动计算批次大小:")
                    # self.signals.update_progress.emit(f"- 模型上下文大小: {model_token_limit} tokens")
                    # self.signals.update_progress.emit(f"- 安全token限制: {safe_token_limit} tokens (带安全边界)")
                    # self.signals.update_progress.emit(f"- 样本{sample_size}行平均每行: {avg_tokens_per_row:.2f} tokens")
                    # self.signals.update_progress.emit(f"- 系统消息占用: {system_tokens} tokens")
                    # self.signals.update_progress.emit(f"- 可用token: {available_tokens} tokens")
                    # self.signals.update_progress.emit(f"- 自动计算的最佳批次大小: {batch_size} 行/批")
                
                num_batches = (len(df) + batch_size - 1) // batch_size
                
                self.signals.update_progress.emit(f"\n数据将被分为 {num_batches} 批处理，每批最多 {batch_size} 行")
                self.signals.update_progress.emit(f"使用模型：{self.model_name}，上下文限制：{model_token_limit} tokens")
                self.signals.update_progress.emit(f"每批数据发送后将等待 {self.wait_time} 秒，避免超出API限制")
                self.signals.update_progress_bar.emit(25)
                
                # 初始化API客户端
                if model_api_type == "moonshot":
                    self.signals.update_progress.emit("\n正在初始化Moonshot API客户端...")
                    client = OpenAI(
                        api_key=self.api_key,
                        base_url="https://api.moonshot.cn/v1"
                    )
                elif model_api_type == "siliconflow" and SILICON_FLOW_AVAILABLE:
                    self.signals.update_progress.emit("\n正在初始化硅基流动 API客户端...")
                    try:
                        # 检查API密钥格式
                        api_key = self.api_key
                        # 如果API密钥以"Bearer "开头，移除这个前缀，因为SiliconFlowAPI会自动添加
                        if api_key.startswith("Bearer "):
                            api_key = api_key[7:]  # 移除"Bearer "前缀
                            self.signals.update_progress.emit("注意: 已移除API密钥中的'Bearer '前缀，因为客户端会自动添加")
                        # 如果API密钥为空，给出明确提示
                        if not api_key or api_key == "此处需要替换为您的硅基流动API密钥":
                            raise ValueError("硅基流动API密钥未设置，请在高级设置中配置有效的API密钥")
                            
                        self.signals.update_progress.emit(f"使用API密钥格式: {api_key[:10]}...{api_key[-4:]}")
                        siliconflow_client = SiliconFlowAPI(api_key=api_key)
                        self.signals.update_progress.emit("硅基流动API客户端初始化成功")
                    except Exception as e:
                        self.signals.error.emit(f"初始化硅基流动API客户端失败: {str(e)}")
                        return
                elif model_api_type == "volcano":
                    self.signals.update_progress.emit("\n正在初始化火山引擎 API客户端...")
                    try:
                        # 初始化火山引擎API客户端
                        volcano_client = OpenAI(
                            api_key=self.api_key,
                            base_url="https://ark.cn-beijing.volces.com/api/v3"
                        )
                        self.signals.update_progress.emit("火山引擎API客户端初始化成功")
                    except Exception as e:
                        self.signals.error.emit(f"初始化火山引擎API客户端失败: {str(e)}")
                        return
                else:
                    self.signals.error.emit(f"错误：不支持的API类型 {model_api_type} 或缺少必要的库")
                    return
                
                # 设置使用的模型
                model_name = self.model_name
                self.signals.update_progress.emit(f"\n使用模型：{model_name}")
                
                # 分批发送数据并等待
                for i in range(num_batches):
                    if self.stop_requested:
                        self.signals.update_progress.emit("分析任务已被用户取消")
                        return
                        
                    start_idx = i * batch_size
                    end_idx = min((i + 1) * batch_size, len(df))
                    
                    batch_df = df.iloc[start_idx:end_idx]
                    
                    # 转换为文本格式
                    batch_text = f"\n数据批次 {i+1}/{num_batches} (行 {start_idx+1} 到 {end_idx})：\n"
                    batch_text += batch_df.to_string(max_rows=None, max_cols=None)
                    
                    # 计算文字数量和估算的tokens数量
                    chars_count = len(batch_text)
                    tokens_estimate = self.num_tokens_from_string(batch_text)
                    
                    # 注释掉批次发送的详细日志
                    # 注释掉批次信息的详细日志
                    # self.signals.update_progress.emit(f"\n准备发送数据批次 {i+1}/{num_batches} (行 {start_idx+1} 到 {end_idx})")
                    # self.signals.update_progress.emit(f"批次大小: {chars_count} 字符, 估算 {tokens_estimate} tokens")
                    
                    # 更新进度条
                    progress_percent = 25 + (i / num_batches) * 50
                    self.signals.update_progress_bar.emit(int(progress_percent))
                    
                    # 发送当前批次数据
                    try:
                        self.signals.update_progress.emit(f"正在处理数据批次 {i+1}/{num_batches}...")
                        
                        # 构建当前批次的消息 - 只包含系统消息和当前批次，不包含历史
                        current_messages = initial_system_messages.copy()
                        current_messages.append({
                            "role": "user", 
                            "content": f"这是数据批次 {i+1}/{num_batches}，总共{num_batches}批:\n{batch_text}\n\n请只回复确认收到即可，例如'已接收数据批次{i+1}/{num_batches}'，不需要分析。"
                        })
                        
                        # 注释掉请求总大小的日志
                        # total_chars = sum(len(msg["content"]) for msg in current_messages)
                        # total_tokens = self.num_tokens_from_string(json.dumps([msg["content"] for msg in current_messages], ensure_ascii=False))
                        # self.signals.update_progress.emit(f"当前请求总大小: {total_chars} 字符, 估算 {total_tokens} tokens")
                        
                        # 根据API类型不同，发送请求
                        if model_api_type == "moonshot":
                            # 发送请求到Moonshot API
                            response = client.chat.completions.create(
                                model=model_name,
                                messages=current_messages,
                                temperature=0.2
                            )
                            
                            # 获取回复
                            batch_response = response.choices[0].message.content
                            
                            # 如果API返回了token使用信息，显示出来
                            if hasattr(response, 'usage') and response.usage:
                                self.signals.update_progress.emit(f"API报告的token使用情况:")
                                self.signals.update_progress.emit(f"- 输入tokens: {response.usage.prompt_tokens}")
                                self.signals.update_progress.emit(f"- 输出tokens: {response.usage.completion_tokens}")
                                self.signals.update_progress.emit(f"- 总tokens: {response.usage.total_tokens}")
                                
                        elif model_api_type == "siliconflow":
                            # 发送请求到硅基流动API
                            try:
                                # 注释掉API详细信息的日志
                                # self.signals.update_progress.emit(f"正在发送批次 {i+1}/{num_batches} 到硅基流动API [模型: {model_name}]...")
                                
                                # 尝试使用简化版API客户端，避免400错误
                                try:
                                    # 注释掉API客户端详细信息的日志
                                    # self.signals.update_progress.emit("尝试使用简化版API客户端发送请求...")
                                    simplified_client = SimplifiedSiliconFlowAPI(api_key=api_key)
                                    
                                    # 发送请求
                                    response = simplified_client.chat_completion(
                                        messages=current_messages,
                                        model=model_name,
                                        temperature=0.2,
                                        max_tokens=500
                                    )
                                    
                                    # 注释掉API请求成功的详细日志
                                    # self.signals.update_progress.emit("简化版API请求成功!")
                                except Exception as simplified_error:
                                    # 如果简化版失败，尝试原始版本
                                    self.signals.update_progress.emit(f"API请求失败: {str(simplified_error)}，尝试使用备用方法...")
                                
                                # 获取回复
                                batch_response = response.get("choices", [{}])[0].get("message", {}).get("content", "")
                                
                                # 打印API调用结果基本信息
                                self.signals.update_progress.emit("硅基流动API请求成功")
                                
                                # 如果API返回了token使用信息，显示出来
                                if "usage" in response:
                                    usage = response["usage"]
                                    self.signals.update_progress.emit(f"API报告的token使用情况:")
                                    self.signals.update_progress.emit(f"- 输入tokens: {usage.get('prompt_tokens', 'N/A')}")
                                    self.signals.update_progress.emit(f"- 输出tokens: {usage.get('completion_tokens', 'N/A')}")
                                    self.signals.update_progress.emit(f"- 总tokens: {usage.get('total_tokens', 'N/A')}")
                            except Exception as e:
                                error_msg = str(e)
                                self.signals.update_progress.emit(f"硅基流动API请求失败: {error_msg}")
                                
                                # 提供更详细的硅基流动API错误信息和建议
                                if "401" in error_msg and "Unauthorized" in error_msg:
                                    self.signals.update_progress.emit("\n[API密钥错误] 401 Unauthorized表示认证失败，可能原因:")
                                    self.signals.update_progress.emit("1. API密钥格式不正确 - 应为纯密钥格式，如'sk-xxx'")
                                    self.signals.update_progress.emit("2. API密钥已过期或被禁用")
                                    self.signals.update_progress.emit("3. API密钥权限不足，无法访问请求的资源或模型")
                                    self.signals.update_progress.emit("\n请在高级设置中检查并更新您的API密钥")
                                
                                # 重新抛出异常，让外层错误处理捕获
                                raise
                        
                        elif model_api_type == "volcano":
                            # 发送请求到火山引擎API
                            try:
                                # 发送请求
                                response = volcano_client.chat.completions.create(
                                    model=model_name,
                                    messages=current_messages,
                                    temperature=0.2
                                )
                                
                                # 获取回复
                                batch_response = response.choices[0].message.content
                                
                                # 打印API调用结果基本信息
                                self.signals.update_progress.emit("火山引擎API请求成功")
                                
                                # 如果API返回了token使用信息，显示出来
                                if hasattr(response, 'usage') and response.usage:
                                    self.signals.update_progress.emit(f"API报告的token使用情况:")
                                    self.signals.update_progress.emit(f"- 输入tokens: {response.usage.prompt_tokens}")
                                    self.signals.update_progress.emit(f"- 输出tokens: {response.usage.completion_tokens}")
                                    self.signals.update_progress.emit(f"- 总tokens: {response.usage.total_tokens}")
                            except Exception as e:
                                error_msg = str(e)
                                self.signals.update_progress.emit(f"火山引擎API请求失败: {error_msg}")
                                raise
                        
                        self.signals.update_progress.emit(f"模型回复: {batch_response}")
                        
                        # 如果不是最后一批，等待指定时间再继续
                        if i < num_batches - 1:
                            self.signals.update_progress.emit(f"等待{self.wait_time}秒后发送下一批数据...")
                            
                            for j in range(self.wait_time, 0, -1):
                                if self.stop_requested:
                                    self.signals.update_progress.emit("分析任务已被用户取消")
                                    return
                                # 注释掉等待倒计时的日志
                                # self.signals.update_progress.emit(f"等待中... {j} 秒")
                                time.sleep(1)
                                
                            # self.signals.update_progress.emit("等待完成，继续发送...")
                        
                    except Exception as e:
                        error_str = str(e)
                        
                        # 打印详细的错误信息
                        self.signals.update_progress.emit(f"\n发送批次 {i+1} 时出错:")
                        self.signals.update_progress.emit(f"错误详情: {error_str}")
                        self.signals.update_progress.emit(f"错误类型: {type(e).__name__}")
                        
                        # 如果是API错误，尝试提取更多信息
                        if hasattr(e, 'response'):
                            try:
                                error_response = e.response
                                self.signals.update_progress.emit(f"API错误状态码: {error_response.status_code}")
                                error_json = error_response.json()
                                self.signals.update_progress.emit(f"API错误信息: {json.dumps(error_json, ensure_ascii=False, indent=2)}")
                            except:
                                self.signals.update_progress.emit("无法解析API错误响应")
                        
                        # 检查是否是频率限制错误
                        if "rate_limit_reached_error" in error_str:
                            retry_wait = self.wait_time + 15  # 额外等待15秒
                            self.signals.update_progress.emit(f"调用API时遇到频率限制，等待{retry_wait}秒后重试...")
                            
                            # 倒计时显示
                            for j in range(retry_wait, 0, -1):
                                if self.stop_requested:
                                    self.signals.update_progress.emit("分析任务已被用户取消")
                                    return
                                # 注释掉等待倒计时的日志
                                # self.signals.update_progress.emit(f"等待中... {j} 秒")
                                time.sleep(1)
                            # self.signals.update_progress.emit("等待完成，重试发送...")
                            
                            # 重试当前批次
                            i -= 1  # 回退一步，重试当前批次
                            continue
                        
                        # 检查是否是token超限错误
                        if "exceeded model token limit" in error_str or "max_total_tokens" in error_str or "length of prompt_tokens" in error_str:
                            # 更激进地减小批次大小
                            current_tokens_estimate = tokens_estimate  # 当前批次的估计token数
                            
                            # 对于Qwen模型，使用更激进的减小策略
                            if "qwen" in self.model_name.lower():
                                reduced_batch_size = max(10, int(batch_size * 0.3))  # 减少到原来的30%
                                self.signals.update_progress.emit(f"检测到Qwen模型token限制错误，将使用更激进的批次大小减小策略")
                            else:
                                reduced_batch_size = max(10, int(batch_size * 0.5))  # 直接减半，更保守的策略
                            
                            self.signals.update_progress.emit(f"遇到token超限错误，将批次大小从{batch_size}减少到{reduced_batch_size}行，重新尝试...")
                            self.signals.update_progress.emit(f"当前批次估计token数: {current_tokens_estimate}, 超过了模型限制")
                            self.signals.update_progress.emit(f"错误详情: {error_str}")
                            
                            # 调整参数并重启处理
                            batch_size = reduced_batch_size
                            
                            # 重新计算批次数
                            num_batches = (len(df) + batch_size - 1) // batch_size
                            self.signals.update_progress.emit(f"数据将重新分为 {num_batches} 批处理，每批最多 {batch_size} 行")
                            
                            # 重新开始处理
                            self.signals.update_progress.emit("重新开始批次处理...")
                            i = -1  # 下一次循环会从0开始
                            continue
                        
                        # 其他错误，终止程序
                        self.signals.error.emit(f"发送数据时出错: {error_str}")
                        return
                
                # 所有批次发送完毕，发送最终分析请求
                self.signals.update_progress.emit("\n所有数据批次已发送完毕，现在请求模型进行分析...")
                self.signals.update_progress_bar.emit(75)
                
                # 为最终分析创建新的对话历史，只包含必要信息，不包含原始数据
                final_conversation_history = [
                    {"role": "system", "content": """你是一位数据分析助手。
你需要分析一个Excel数据集中的异常值。

数据集信息已经提供给你，你的任务是根据数据特征识别可能的异常值和数据质量问题。
请提供详细的分析，包括异常值的位置、类型和可能的原因，以及如何修正这些问题。
分析应该清晰、结构化，帮助用户理解数据中的问题并采取行动。"""},
                    {"role": "system", "content": data_info},
                    {"role": "user", "content": f"""此Excel数据集共有{len(df)}行和{len(df.columns)}列。
我已向你提供了该数据集的列名和数据类型信息。

现在请详细分析这个Excel数据集中哪些字段的值存在异常。

请提供以下信息：
1. 异常字段的名称
2. 异常值描述（包括出现异常值的具体数据，以及具体数据的异常值描述）
3. 问题分析（为什么这些值被认为是异常的，可能的原因）
4. 处理建议（修正这些异常值的具体方法）

如果没有发现异常，请说明数据集看起来正常。

在回答前，请先简要描述一下这个数据集的基本情况，包括：
- 数据集的行数和列数
- 各列的数据类型
- 数据集的主要内容和用途
- 数据的时间范围（如果适用）

"""}
                ]
                
                # 计算最终请求的总tokens估算
                final_request_chars = sum(len(msg["content"]) for msg in final_conversation_history)
                final_request_tokens = self.num_tokens_from_string(json.dumps([msg["content"] for msg in final_conversation_history], ensure_ascii=False))
                self.signals.update_progress.emit(f"最终分析请求总大小: {final_request_chars} 字符, 估算 {final_request_tokens} tokens")
                
                # 发送聊天请求
                self.signals.update_progress.emit("\n正在请求模型分析数据集...")
                self.signals.update_progress.emit(f"(使用{model_name}模型分析，过程可能需要几分钟时间，请耐心等待)")
                self.signals.update_progress_bar.emit(80)
                
                start_time = time.time()
                
                # 添加重试机制
                max_retries = 3
                retry_count = 0
                retry_delay = 45
                
                analysis_result = None
                
                while retry_count < max_retries and not self.stop_requested:
                    try:
                        self.signals.update_progress.emit(f"正在发送分析请求... (尝试 {retry_count + 1}/{max_retries})")
                        
                        # 根据API类型不同，发送请求
                        if model_api_type == "moonshot":
                            # 使用精简的对话历史发送请求到Moonshot API
                            response = client.chat.completions.create(
                                model=model_name,
                                messages=final_conversation_history,
                                temperature=0.2
                            )
                            
                            # 获取分析结果
                            analysis_result = response.choices[0].message.content
                            
                            # 如果API返回了token使用信息，显示出来
                            if hasattr(response, 'usage') and response.usage:
                                self.signals.update_progress.emit(f"API报告的最终分析token使用情况:")
                                self.signals.update_progress.emit(f"- 输入tokens: {response.usage.prompt_tokens}")
                                self.signals.update_progress.emit(f"- 输出tokens: {response.usage.completion_tokens}")
                                self.signals.update_progress.emit(f"- 总tokens: {response.usage.total_tokens}")
                                
                        elif model_api_type == "siliconflow":
                            # 发送请求到硅基流动API
                            try:
                                # 打印API请求信息
                                self.signals.update_progress.emit(f"正在发送最终分析请求到硅基流动API [模型: {model_name}]...")
                                
                                # 尝试使用简化版API客户端，避免400错误
                                try:
                                    # 初始化简化版API客户端
                                    self.signals.update_progress.emit("尝试使用简化版API客户端发送最终分析请求...")
                                    simplified_client = SimplifiedSiliconFlowAPI(api_key=api_key)
                                    
                                    # 发送请求
                                    response = simplified_client.chat_completion(
                                        messages=final_conversation_history,
                                        model=model_name,
                                        temperature=0.2,
                                        max_tokens=2000 if "qwen" in model_name.lower() else 4000  # 对Qwen模型使用更保守的max_tokens设置
                                    )
                                    
                                    self.signals.update_progress.emit("简化版API请求成功!")
                                except Exception as simplified_error:
                                    # 如果简化版失败，尝试原始版本
                                    self.signals.update_progress.emit(f"简化版API请求失败: {str(simplified_error)}，尝试使用原始版本...")
                                    
                                    try:
                                        # 发送请求
                                        response = siliconflow_client.chat_completion(
                                            messages=final_conversation_history,
                                            model=model_name,
                                            temperature=0.2,
                                            max_tokens=4000
                                        )
                                        self.signals.update_progress.emit("原始版API请求成功!")
                                    except Exception as original_error:
                                        # 如果原始版也失败，尝试直接API请求
                                        self.signals.update_progress.emit(f"原始版API请求失败: {str(original_error)}，尝试使用直接API请求...")
                                        
                                        # 使用备用方法
                                        response = self.direct_api_request(
                                            messages=final_conversation_history,
                                            api_key=api_key,
                                            model_name=model_name
                                        )
                                        self.signals.update_progress.emit("直接API请求成功!")
                                
                                # 获取分析结果
                                analysis_result = response.get("choices", [{}])[0].get("message", {}).get("content", "")
                                
                                # 打印API调用结果基本信息
                                self.signals.update_progress.emit("硅基流动API最终分析请求成功")
                                
                                # 如果API返回了token使用信息，显示出来
                                if "usage" in response:
                                    usage = response["usage"]
                                    self.signals.update_progress.emit(f"API报告的最终分析token使用情况:")
                                    self.signals.update_progress.emit(f"- 输入tokens: {usage.get('prompt_tokens', 'N/A')}")
                                    self.signals.update_progress.emit(f"- 输出tokens: {usage.get('completion_tokens', 'N/A')}")
                                    self.signals.update_progress.emit(f"- 总tokens: {usage.get('total_tokens', 'N/A')}")
                            except Exception as e:
                                error_msg = str(e)
                                self.signals.update_progress.emit(f"硅基流动API最终分析请求失败: {error_msg}")
                                
                                # 提供更详细的硅基流动API错误信息和建议
                                if "401" in error_msg and "Unauthorized" in error_msg:
                                    self.signals.update_progress.emit("\n[API密钥错误] 401 Unauthorized表示认证失败，可能原因:")
                                    self.signals.update_progress.emit("1. API密钥格式不正确 - 应为纯密钥格式，如'sk-xxx'")
                                    self.signals.update_progress.emit("2. API密钥已过期或被禁用")
                                    self.signals.update_progress.emit("3. API密钥权限不足，无法访问请求的资源或模型")
                                    self.signals.update_progress.emit("\n请在高级设置中检查并更新您的API密钥")
                                
                                # 重新抛出异常，让外层错误处理捕获
                                raise
                        
                        elif model_api_type == "volcano":
                            # 发送请求到火山引擎API
                            try:
                                # 发送请求
                                response = volcano_client.chat.completions.create(
                                    model=model_name,
                                    messages=final_conversation_history,
                                    temperature=0.2
                                )
                                
                                # 获取分析结果
                                analysis_result = response.choices[0].message.content
                                
                                # 如果API返回了token使用信息，显示出来
                                if hasattr(response, 'usage') and response.usage:
                                    self.signals.update_progress.emit(f"API报告的最终分析token使用情况:")
                                    self.signals.update_progress.emit(f"- 输入tokens: {response.usage.prompt_tokens}")
                                    self.signals.update_progress.emit(f"- 输出tokens: {response.usage.completion_tokens}")
                                    self.signals.update_progress.emit(f"- 总tokens: {response.usage.total_tokens}")
                            except Exception as e:
                                error_msg = str(e)
                                self.signals.update_progress.emit(f"火山引擎API最终分析请求失败: {error_msg}")
                                raise
                        
                        # 计算处理时间
                        process_time = time.time() - start_time
                        
                        self.signals.update_progress_bar.emit(90)
                        
                        # 打印模型的回答
                        self.signals.update_progress.emit("\n" + "="*80 + "\n")
                        self.signals.update_progress.emit("数据集异常分析结果：\n")
                        self.signals.update_progress.emit(analysis_result)
                        self.signals.update_progress.emit("\n" + "="*80)
                        self.signals.update_progress.emit(f"\n分析耗时：{process_time:.2f} 秒")
                        
                        # 保存分析结果到Word文件
                        self.signals.update_progress.emit("\n正在保存分析结果到Word文档...")
                        
                        # 保存为Word文档
                        saved_path = self.save_to_word(
                            analysis_result, 
                            file_path.name, 
                            process_time
                        )
                        
                        self.signals.update_progress.emit(f"\n分析结果已保存到: {saved_path}")
                        self.signals.update_progress_bar.emit(100)
                        
                        # 发送分析完成信号
                        self.signals.analysis_complete.emit(analysis_result, str(saved_path))
                        
                        # 成功获取结果，跳出循环
                        break
                    
                    except Exception as e:
                        error_str = str(e)
                        retry_count += 1
                        
                        # 打印详细的错误信息
                        self.signals.update_progress.emit(f"\n错误详情: {error_str}")
                        self.signals.update_progress.emit(f"错误类型: {type(e).__name__}")
                        
                        # 如果是API错误，尝试提取更多信息
                        if hasattr(e, 'response'):
                            try:
                                error_response = e.response
                                self.signals.update_progress.emit(f"API错误状态码: {error_response.status_code}")
                                error_json = error_response.json()
                                self.signals.update_progress.emit(f"API错误信息: {json.dumps(error_json, ensure_ascii=False, indent=2)}")
                            except:
                                self.signals.update_progress.emit("无法解析API错误响应")
                        
                        # 检查是否是频率限制错误
                        if "rate_limit_reached_error" in error_str:
                            # 固定等待45秒
                            wait_time = 45
                            
                            self.signals.update_progress.emit(f"调用API时遇到频率限制，等待 {wait_time} 秒后重试... (尝试 {retry_count}/{max_retries})")
                            
                            # 倒计时显示
                            for i in range(wait_time, 0, -1):
                                if self.stop_requested:
                                    self.signals.update_progress.emit("分析任务已被用户取消")
                                    return
                                # 注释掉等待倒计时的日志
                                # self.signals.update_progress.emit(f"等待中... {i} 秒")
                                time.sleep(1)
                            # self.signals.update_progress.emit("等待完成，继续请求...")
                            continue
                        
                        # 如果不是频率限制错误，或者已经达到最大重试次数，则抛出异常
                        if retry_count >= max_retries:
                            self.signals.error.emit(f"调用API失败: 达到最大重试次数 ({max_retries})")
                            return
                        
                        self.signals.update_progress.emit(f"调用API失败，45秒后重试: {error_str}")
                        # 倒计时显示
                        for i in range(45, 0, -1):
                            if self.stop_requested:
                                self.signals.update_progress.emit("分析任务已被用户取消")
                                return
                            # 注释掉等待倒计时的日志
                            # self.signals.update_progress.emit(f"等待中... {i} 秒")
                            time.sleep(1)
                
                # 如果所有重试都失败，抛出异常
                if retry_count >= max_retries:
                    self.signals.error.emit("调用API失败: 达到最大重试次数")
                
                # 如果用户请求停止
                if self.stop_requested:
                    self.signals.update_progress.emit("分析任务已被用户取消")
                    
            except pd.errors.EmptyDataError:
                self.signals.error.emit("错误：Excel文件为空")
            except pd.errors.ParserError:
                self.signals.error.emit("错误：无法解析Excel文件，文件可能已损坏")
            except Exception as e:
                self.signals.error.emit(f"分析过程中出错：{str(e)}")
        
        except Exception as e:
            self.signals.error.emit(f"程序运行出错：{str(e)}")
        
        finally:
            # 无论成功与否，确保UI状态被重置
            if not self.stop_requested:
                self.signals.update_progress.emit("分析任务已完成")
    
    def stop(self):
        """请求停止分析"""
        self.stop_requested = True
        self.signals.update_progress.emit("正在尝试停止分析任务...")

class ExcelAnalyzerApp(QMainWindow):
    """Excel数据异常分析GUI应用"""
    def __init__(self):
        super().__init__()
        
        # 设置窗口标题和大小
        self.setWindowTitle("Excel数据集异常值分析工具")
        self.setMinimumSize(1000, 700)
        
        # 应用自定义样式
        self.apply_custom_style()
        
        # 初始化UI
        self.init_ui()
        
        # 分析工作线程
        self.worker = None
        
        # 结果文件路径
        self.result_path = ""
        
        # 结果展开窗口
        self.expanded_result_dialog = None
    
    def apply_custom_style(self):
        """应用自定义样式"""
        # 设置应用程序样式表
        self.setStyleSheet(f"""
            QMainWindow, QDialog {{
                background-color: white;
            }}
            QGroupBox {{
                border: 1px solid #d0d0d0;
                border-radius: 5px;
                margin-top: 10px;
                font-weight: bold;
                color: rgb(60, 60, 60);
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 3px;
                color: rgb(60, 60, 60);
            }}
            QPushButton {{
                background-color: rgb(158, 33, 65);
                color: white;
                border: none;
                border-radius: 4px;
                padding: 6px 12px;
                min-height: 30px;
            }}
            QPushButton:hover {{
                background-color: rgb(178, 53, 85);
            }}
            QPushButton:pressed {{
                background-color: rgb(138, 13, 45);
            }}
            QPushButton:disabled {{
                background-color: rgb(166, 166, 166);
            }}
            QProgressBar {{
                border: 1px solid #d0d0d0;
                border-radius: 4px;
                text-align: center;
                color: white;
                font-weight: bold;
            }}
            QProgressBar::chunk {{
                background-color: rgb(177, 136, 101);
            }}
            QLineEdit, QComboBox, QSpinBox {{
                border: 1px solid #d0d0d0;
                border-radius: 4px;
                padding: 4px;
                min-height: 24px;
                color: rgb(60, 60, 60);
                background-color: white;
            }}
            QComboBox::drop-down {{
                border: none;
                background-color: white;
                width: 20px;
            }}
            QComboBox::down-arrow {{
                width: 14px;
                height: 14px;
                image: url(data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHZpZXdCb3g9IjAgMCAyNCAyNCI+PHBhdGggZD0iTTcgMTBsNSA1IDUtNXoiIGZpbGw9IiM2NjY2NjYiLz48L3N2Zz4=);
            }}
            QComboBox QAbstractItemView {{
                border: 1px solid #d0d0d0;
                border-radius: 0px;
                background-color: white;
                selection-background-color: #f0f0f0;
                selection-color: rgb(60, 60, 60);
            }}
            QTabWidget::pane {{
                border: 1px solid #d0d0d0;
                border-radius: 4px;
            }}
            QTabBar::tab {{
                background-color: #f0f0f0;
                border: 1px solid #d0d0d0;
                border-bottom-color: none;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                padding: 6px 12px;
                margin-right: 2px;
                color: rgb(80, 80, 80);
            }}
            QTabBar::tab:selected {{
                background-color: white;
                border-bottom-color: white;
                color: rgb(158, 33, 65);
                font-weight: bold;
            }}
            QTabBar::tab:hover {{
                background-color: #e0e0e0;
            }}
            QLabel {{
                color: rgb(60, 60, 60);
            }}
            QTextEdit {{
                color: rgb(60, 60, 60);
                background-color: white;
                selection-background-color: rgb(177, 136, 101);
                selection-color: white;
            }}
        """)
    
    def init_ui(self):
        """初始化用户界面"""
        # 创建中央窗口部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # 设置主布局
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        main_layout.setSpacing(15)
        
        # 创建标题标签
        title_label = QLabel("Excel数据集异常值分析工具")
        title_label.setAlignment(Qt.AlignCenter)
        title_font = QFont()
        title_font.setPointSize(18)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_palette = title_label.palette()
        title_palette.setColor(QPalette.WindowText, COLOR_BURGUNDY)
        title_label.setPalette(title_palette)
        main_layout.addWidget(title_label)
        
        # 创建子标题标签
        subtitle_label = QLabel("使用大模型分析Excel数据集中的异常值")
        subtitle_label.setAlignment(Qt.AlignCenter)
        subtitle_font = QFont()
        subtitle_font.setPointSize(12)
        subtitle_label.setFont(subtitle_font)
        subtitle_palette = subtitle_label.palette()
        subtitle_palette.setColor(QPalette.WindowText, COLOR_GOLD)
        subtitle_label.setPalette(subtitle_palette)
        main_layout.addWidget(subtitle_label)
        
        # 添加分隔线
        line = QFrame()
        line.setFrameShape(QFrame.HLine)
        line.setFrameShadow(QFrame.Sunken)
        line_palette = line.palette()
        line_palette.setColor(QPalette.WindowText, COLOR_GRAY)
        line.setPalette(line_palette)
        main_layout.addWidget(line)
        
        # 文件选择区域
        file_group = QGroupBox("选择Excel文件 (最大500KB)")
        file_layout = QHBoxLayout()
        file_layout.setContentsMargins(15, 15, 15, 15)
        
        self.file_path_edit = QLineEdit()
        self.file_path_edit.setPlaceholderText("请选择要分析的Excel文件...")
        self.file_path_edit.setReadOnly(True)
        
        # 添加文件图标到浏览按钮
        browse_button = QPushButton("浏览...")
        browse_button.setIcon(self.style().standardIcon(getattr(QStyle, "SP_DialogOpenButton", "SP_FileDialogStart")))
        browse_button.setFixedWidth(100)
        browse_button.clicked.connect(self.browse_file)
        
        file_layout.addWidget(self.file_path_edit, 3)
        file_layout.addWidget(browse_button, 1)
        
        file_group.setLayout(file_layout)
        main_layout.addWidget(file_group)
        
        # 参数设置区域 - 只保留模型选择
        params_group = QGroupBox("分析参数设置")
        params_layout = QFormLayout()
        params_layout.setContentsMargins(15, 15, 15, 15)
        
        # API密钥设置 - 隐藏进高级设置
        model_config = MODEL_CONFIG.get(DEFAULT_MODEL, {})
        default_api_key = model_config.get("api_key", DEFAULT_MOONSHOT_API_KEY)
        self.api_key_edit = QLineEdit(default_api_key)
        self.api_key_edit.setEchoMode(QLineEdit.Password)
        self.api_key_edit.hide()  # 隐藏API密钥输入框
        
        # 模型选择
        self.model_combo = QComboBox()
        # 添加所有支持的模型，不显示图标，只显示模型名
        for model_name, config in MODEL_CONFIG.items():
            # 移除图标，只使用模型名
            self.model_combo.addItem(f"{model_name}", model_name)
        
        # 设置下拉菜单的最小宽度，确保能够显示完整的模型名称
        self.model_combo.setMinimumWidth(300)
        
        # 设置当前选中的模型
        for i in range(self.model_combo.count()):
            if self.model_combo.itemData(i) == DEFAULT_MODEL:
                self.model_combo.setCurrentIndex(i)
                break
        
        # 当模型变化时，调整批次大小建议和API key
        self.model_combo.currentIndexChanged.connect(self.update_model_settings)
        
        # 设置固定的其他参数（隐藏）
        self.batch_size_spin = QSpinBox()
        self.batch_size_spin.setValue(0)  # 默认使用自动计算批次大小
        self.batch_size_spin.hide()
        
        self.wait_time_spin = QSpinBox()
        self.wait_time_spin.setValue(DEFAULT_WAIT_TIME)
        self.wait_time_spin.hide()
        
        self.max_rows_spin = QSpinBox()
        self.max_rows_spin.setValue(0)
        self.max_rows_spin.hide()
        
        params_layout.addRow("使用模型:", self.model_combo)
        
        # 调整表单布局的字段比例，给下拉菜单更多空间
        params_layout.setLabelAlignment(Qt.AlignLeft | Qt.AlignVCenter)
        params_layout.setFieldGrowthPolicy(QFormLayout.ExpandingFieldsGrow)
        
        # 注释掉高级设置按钮
        # self.advanced_button = QPushButton("高级设置")
        # self.advanced_button.setIcon(self.style().standardIcon(getattr(QStyle, "SP_FileDialogDetailedView", "SP_FileDialogListView")))
        # self.advanced_button.setFixedWidth(120)
        # self.advanced_button.clicked.connect(self.toggle_advanced_settings)
        # params_layout.addRow("", self.advanced_button)
        
        params_group.setLayout(params_layout)
        main_layout.addWidget(params_group)
        
        # 操作按钮
        buttons_layout = QHBoxLayout()
        
        self.analyze_button = QPushButton("AI一键分析")
        self.analyze_button.setMinimumHeight(45)
        self.analyze_button.setIcon(self.style().standardIcon(getattr(QStyle, "SP_CommandLink", "SP_ArrowRight")))
        self.analyze_button.setIconSize(QSize(20, 20))
        font = QFont()
        font.setBold(True)
        font.setPointSize(12)
        self.analyze_button.setFont(font)
        self.analyze_button.clicked.connect(self.start_analysis)
        
        self.stop_button = QPushButton("停止分析")
        self.stop_button.setIcon(self.style().standardIcon(getattr(QStyle, "SP_DialogCancelButton", "SP_BrowserStop")))
        self.stop_button.setEnabled(False)
        self.stop_button.clicked.connect(self.stop_analysis)
        
        buttons_layout.addWidget(self.analyze_button)
        buttons_layout.addWidget(self.stop_button)
        
        main_layout.addLayout(buttons_layout)
        
        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setFixedHeight(24)
        main_layout.addWidget(self.progress_bar)
        
        # 创建选项卡窗口
        self.tab_widget = QTabWidget()
        
        # 进度日志
        self.progress_log = QTextEdit()
        self.progress_log.setReadOnly(True)
        
        # 分析结果
        self.result_container = QWidget()
        result_layout = QVBoxLayout(self.result_container)
        
        self.result_text = QTextEdit()
        self.result_text.setReadOnly(True)
        result_layout.addWidget(self.result_text)
        
        # 添加展开按钮
        expand_button = QPushButton("展开查看完整结果")
        expand_button.setIcon(self.style().standardIcon(getattr(QStyle, "SP_FileDialogContentsView", "SP_FileDialogInfoView")))
        expand_button.clicked.connect(self.show_expanded_result)
        result_layout.addWidget(expand_button)
        
        # 添加选项卡
        self.tab_widget.addTab(self.progress_log, "处理日志")
        self.tab_widget.addTab(self.result_container, "分析结果")
        
        main_layout.addWidget(self.tab_widget, 1)
        
        # 底部状态栏布局
        status_layout = QHBoxLayout()
        
        self.status_label = QLabel("就绪")
        status_layout.addWidget(self.status_label, 1)
        
        self.open_result_button = QPushButton("打开结果文件")
        self.open_result_button.setIcon(self.style().standardIcon(getattr(QStyle, "SP_FileIcon", "SP_DialogOpenButton")))
        self.open_result_button.setEnabled(False)
        self.open_result_button.clicked.connect(self.open_result_file)
        
        status_layout.addWidget(self.open_result_button, 1)
        
        main_layout.addLayout(status_layout)
    
    def toggle_api_key_visibility(self, checked):
        """切换API密钥的可见性"""
        self.api_key_edit.setEchoMode(QLineEdit.Normal if checked else QLineEdit.Password)
    
    def toggle_advanced_settings(self):
        """高级设置已被移除的提示"""
        QMessageBox.information(
            self, 
            "功能已简化", 
            "程序已配置为使用最佳默认设置，无需手动配置。\n\n"
            "批次大小: 自动计算\n"
            "处理模式: 使用模型最佳参数"
        )
    
    def start_analysis(self):
        """开始分析"""
        if not self.file_path_edit.text():
            QMessageBox.warning(self, "警告", "请先选择Excel文件")
            return
        
        # 获取当前选择的模型信息
        model_name = self.model_combo.currentData()
        model_config = MODEL_CONFIG.get(model_name, {})
        api_key = self.api_key_edit.text()
        
        # 使用自动计算的批次大小
        batch_size = 0  # 自动计算
        wait_time = DEFAULT_WAIT_TIME
        max_rows = 0  # 不限制行数
            
        self.worker = AnalysisWorker(
            self.file_path_edit.text(),
            api_key,
            batch_size,
            wait_time,
            model_name,
            max_rows
        )
        self.worker.signals.update_progress.connect(self.update_progress)
        self.worker.signals.update_progress_bar.connect(self.update_progress_bar)
        self.worker.signals.analysis_complete.connect(self.analysis_complete)
        self.worker.signals.error.connect(self.handle_error)
        self.worker.start()
        self.analyze_button.setEnabled(False)
        self.stop_button.setEnabled(True)
        
        # 自动切换到处理日志选项卡
        self.tab_widget.setCurrentIndex(0)
    
    def stop_analysis(self):
        """停止分析"""
        if self.worker:
            self.worker.stop()
            self.worker = None
            self.analyze_button.setEnabled(True)
            self.stop_button.setEnabled(False)
    
    def update_progress(self, message):
        """更新进度信息"""
        self.progress_log.append(message)
    
    def update_progress_bar(self, value):
        """更新进度条"""
        self.progress_bar.setValue(value)
    
    def analysis_complete(self, result, saved_path):
        """分析完成"""
        self.result_text.setText(result)
        self.result_path = saved_path  # 保存结果文件路径
        self.open_result_button.setEnabled(True)
        self.analyze_button.setEnabled(True)
        self.stop_button.setEnabled(False)
    
    def handle_error(self, error_message):
        """处理错误"""
        self.progress_log.append(f"错误: {error_message}")
        self.open_result_button.setEnabled(False)
    
    def open_result_file(self):
        """打开结果文件"""
        if self.result_path and os.path.exists(self.result_path):
            # 使用platform模块检测操作系统
            import platform
            if platform.system() == 'Windows':
                os.startfile(self.result_path)
            elif platform.system() == 'Darwin':  # macOS
                import subprocess
                subprocess.run(['open', self.result_path], check=True)
            else:  # Linux或其他系统
                import subprocess
                subprocess.run(['xdg-open', self.result_path], check=True)
        else:
            self.progress_log.append("错误：结果文件不存在或尚未生成")
    
    def update_model_settings(self):
        """根据选择的模型更新批次大小建议和API密钥"""
        model_name = self.model_combo.currentData()
        
        if model_name in MODEL_CONFIG:
            config = MODEL_CONFIG[model_name]
            # 获取API类型
            api_type = config.get("api_type", "moonshot")
            
            # 更新API密钥
            api_key = config.get("api_key", "")
            self.api_key_edit.setText(api_key)
            
            # 自动设置为自动计算批次大小
            self.batch_size_spin.setValue(0)
            
            # 如果选择的是硅基流动模型且API密钥为空，提示用户配置
            if api_type == "siliconflow" and (not api_key or api_key == ""):
                QMessageBox.information(
                    self, 
                    "需要设置API密钥", 
                    "您选择了硅基流动API模型，但未找到API密钥配置。\n\n"
                    "程序将使用默认密钥，如需自定义请联系开发者。"
                )
    
    def show_expanded_result(self):
        """显示展开的结果窗口"""
        if not self.expanded_result_dialog:
            self.expanded_result_dialog = QDialog(self)
            self.expanded_result_dialog.setWindowTitle("完整分析结果")
            self.expanded_result_dialog.setMinimumSize(800, 600)
            
            layout = QVBoxLayout(self.expanded_result_dialog)
            
            # 创建可滚动区域和文本编辑器
            scroll_area = QScrollArea()
            scroll_area.setWidgetResizable(True)
            
            expanded_result_text = QTextEdit()
            expanded_result_text.setReadOnly(True)
            expanded_result_text.setHtml(self.result_text.toHtml())
            expanded_result_text.setLineWrapMode(QTextEdit.WidgetWidth)
            
            # 设置字体和样式
            font = QFont()
            font.setPointSize(12)
            expanded_result_text.setFont(font)
            
            # 确保文本颜色可见
            expanded_result_text.setStyleSheet("color: rgb(60, 60, 60); background-color: white;")
            
            scroll_area.setWidget(expanded_result_text)
            layout.addWidget(scroll_area)
            
            # 添加关闭按钮
            close_button = QPushButton("关闭")
            close_button.clicked.connect(self.expanded_result_dialog.close)
            layout.addWidget(close_button)
            
            self.expanded_result_dialog.setLayout(layout)
        else:
            # 更新现有对话框中的文本内容
            text_edit = self.expanded_result_dialog.findChild(QTextEdit)
            if text_edit:
                text_edit.setHtml(self.result_text.toHtml())
        
        self.expanded_result_dialog.show()
    
    def browse_file(self):
        """打开文件选择对话框"""
        file_path, _ = QFileDialog.getOpenFileName(self, "选择Excel文件", "", "Excel Files (*.xlsx);;All Files (*)")
        if file_path:
            # 检查文件大小是否超过500KB
            file_size = os.path.getsize(file_path) / 1024  # 转换为KB
            if file_size > 500:
                QMessageBox.warning(
                    self, 
                    "文件过大", 
                    f"选择的文件大小为 {file_size:.2f} KB，超过了500KB的限制。\n\n"
                    "请选择更小的文件，或者将大文件拆分为多个小文件后再分析。"
                )
                return
            
            # 文件大小符合要求，设置文件路径
            self.file_path_edit.setText(file_path)

if __name__ == "__main__":
    # 检查并安装必要的依赖
    try:
        import docx
    except ImportError:
        print("正在安装python-docx库，用于生成Word文档...")
        import subprocess
        import sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        print("安装完成！")
    
    # 创建并运行应用
    app = QApplication(sys.argv)
    window = ExcelAnalyzerApp()
    window.show()
    sys.exit(app.exec_())
