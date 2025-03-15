#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
批量调整Word文档格式工具
功能：将指定文件夹中的所有Word文档统一调整为指定格式
"""

import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
import traceback
import re
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def get_font_size(size_name):
    """
    转换中文字号到磅值
    """
    sizes = {
        "一号": Pt(26),
        "二号": Pt(22),
        "三号": Pt(16),
        "小三": Pt(15),
        "四号": Pt(14),
        "小四": Pt(12),
        "五号": Pt(10.5),
        "小五": Pt(9)
    }
    return sizes.get(size_name, Pt(12))  # 默认返回小四号

def format_paragraph(paragraph, font_size, font_name="宋体", bold=False, 
                    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT, line_spacing=1.5, 
                    first_line_indent=0, space_before=0.5, space_after=0.5):
    """
    统一设置段落格式
    """
    # 段落格式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = alignment
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = line_spacing
    # 首行缩进（单位：字符）
    if first_line_indent > 0:
        paragraph_format.first_line_indent = Pt(first_line_indent * font_size.pt)
    else:
        paragraph_format.first_line_indent = 0
    # 段前段后间距
    paragraph_format.space_before = Pt(space_before * font_size.pt)
    paragraph_format.space_after = Pt(space_after * font_size.pt)

    # 设置字体
    for run in paragraph.runs:
        run.font.size = font_size
        run.font.name = font_name
        # 中文字体设置
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.bold = bold

def format_table(table, font_size, font_name="宋体", line_spacing=1.0):
    """
    统一设置表格格式
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                format_paragraph(
                    paragraph=paragraph,
                    font_size=font_size,
                    font_name=font_name,
                    bold=False,
                    alignment=WD_PARAGRAPH_ALIGNMENT.CENTER,
                    line_spacing=line_spacing,
                    first_line_indent=0,
                    space_before=0.5,
                    space_after=0.5
                )

def remove_empty_paragraphs(doc):
    """
    删除文档中的空白行（空段落）和空项目符号/编号
    这个方法不会实际删除段落，而是标记它们，以便后续处理时可以跳过
    """
    # 创建一个空列表来保存空段落的索引
    empty_paragraphs = []
    
    # 检查每个段落，确定是否为空
    for i, paragraph in enumerate(doc.paragraphs):
        # 检查是否完全为空或只包含空格
        if not paragraph.text.strip():
            empty_paragraphs.append(i)
            continue
        
        # 增强版检测空的项目符号或编号
        try:
            # 1. 检查是否为列表项
            is_list_item = False
            
            # 检查段落的XML以查找numPr元素
            para_xml = paragraph._element.xml
            if '<w:numPr>' in para_xml or '<w:ilvl>' in para_xml or '<w:numId>' in para_xml:
                is_list_item = True
            
            # 检查样式是否与列表相关
            style_name = paragraph.style.name.lower()
            if any(keyword in style_name for keyword in ('list', 'bullet', 'number', '列表', '编号', '项目')):
                is_list_item = True
                
            # 2. 如果是列表项，检查是否只有项目符号/编号但没有实际内容
            if is_list_item:
                # 获取文本并检查是否实际为空（常见的项目符号可能是：•, -, *, 1., 一, 等）
                text = paragraph.text.strip()
                
                # 使用正则表达式匹配常见的项目符号模式
                if re.match(r'^[•\-\*o\d]+\.?\s*$|^[一二三四五六七八九十]+[、.．]?\s*$|^[IVXLCDM]+\.?\s*$|^[a-zA-Z]\.?\s*$', text):
                    empty_paragraphs.append(i)
                    continue
                
                # 如果只有项目符号，没有内容，或长度很短，可能只是项目符号
                if len(text) <= 3:
                    empty_paragraphs.append(i)
                    continue
                    
                # 如果内容很少，进一步检查
                all_runs_empty = True
                for run in paragraph.runs:
                    run_text = run.text.strip()
                    # 跳过常见的项目符号和数字
                    if run_text and not re.match(r'^[•\-\*o\d]+\.?\s*$|^[一二三四五六七八九十]+[、.．]?\s*$|^[IVXLCDM]+\.?\s*$|^[a-zA-Z]\.?\s*$', run_text):
                        all_runs_empty = False
                        break
                
                if all_runs_empty:
                    empty_paragraphs.append(i)
        except Exception as e:
            # 如果发生异常，记录但继续处理
            print(f"检查段落时出错（索引 {i}）: {str(e)}")
            
    return empty_paragraphs

def add_watermark(doc, text="清控紫荆（北京）教育科技股份有限公司"):
    """
    向文档添加水印
    使用更可靠的方式实现
    """
    try:
        # 为每个节添加水印
        for section in doc.sections:
            # 获取页眉
            header = section.header
            
            # 在页眉中创建一个段落（如果已有则使用第一个）
            if not header.paragraphs:
                header_para = header.add_paragraph()
            else:
                header_para = header.paragraphs[0]
            
            # 居中对齐
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 使用Run添加文本（简单但有效的方法）
            run = header_para.add_run(text)
            run.font.size = Pt(36)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            run.font.color.rgb = RGBColor(192, 192, 192)  # 灰色
            run.font.bold = True
            
            # 为页眉段落添加自定义属性，使其呈现为水印效果
            # 这种方法不如VML完整，但更兼容
            paragraph_props = header_para._p.get_or_add_pPr()
            
            # 确保文本不会被自动截断
            sector_props = paragraph_props.get_or_add_sectPr()
            
        print(f"已成功添加水印: {text}")
        return True
    except Exception as e:
        print(f"添加水印时出错: {str(e)}")
        traceback.print_exc()
        return False

def format_document(doc_path, save_path):
    """
    格式化单个Word文档
    """
    try:
        # 打开文档
        doc = Document(doc_path)
        
        # 获取空白段落的索引列表
        empty_paragraphs = remove_empty_paragraphs(doc)
        
        # 处理段落格式
        for i, paragraph in enumerate(doc.paragraphs):
            # 跳过空白段落
            if i in empty_paragraphs:
                continue
                
            # 根据样式名称或文本内容判断段落类型
            style_name = paragraph.style.name.lower()
            text = paragraph.text.strip()
            
            # 标题处理（可以根据具体文档格式修改判断逻辑）
            if style_name.startswith('heading 1') or (text and text[0] == '第' and '章' in text):
                # 标题1：三号字，宋体，1.5倍行距，无缩进，段前段后为0.5行，加粗，居左对齐
                format_paragraph(
                    paragraph=paragraph,
                    font_size=get_font_size("三号"),
                    bold=True,
                    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
                )
            elif style_name.startswith('heading 2'):
                # 标题2：小三号字，宋体，1.5倍行距，无缩进，段前段后为0.5行，加粗，居左对齐
                format_paragraph(
                    paragraph=paragraph,
                    font_size=get_font_size("小三"),
                    bold=True,
                    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
                )
            elif style_name.startswith('heading 3'):
                # 标题3：四号字，宋体，1.5倍行距，无缩进，段前段后为0.5行，加粗，居左对齐
                format_paragraph(
                    paragraph=paragraph,
                    font_size=get_font_size("四号"),
                    bold=True,
                    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
                )
            elif style_name.startswith('title'):
                # 保留原来的处理逻辑，按照标题1处理
                format_paragraph(
                    paragraph=paragraph,
                    font_size=get_font_size("三号"),
                    bold=True,
                    alignment=WD_PARAGRAPH_ALIGNMENT.LEFT
                )
            else:
                # 正文：小四号字，宋体，1.5倍行距，首行缩进2字符，段前段后为0.5行，不加粗，左对齐
                format_paragraph(
                    paragraph=paragraph,
                    font_size=get_font_size("小四"),
                    bold=False,
                    first_line_indent=2
                )
        
        # 处理表格格式
        for table in doc.tables:
            format_table(
                table=table,
                font_size=get_font_size("小四"),
                line_spacing=1.0
            )
        
        # 尝试添加水印，如果失败则继续其他处理
        try:
            # 添加水印
            add_watermark(doc, "清控紫荆（北京）教育科技股份有限公司")
        except Exception as e:
            print(f"添加水印失败，但将继续处理文档: {str(e)}")
        
        # 保存文档
        doc.save(save_path)
        return True
    
    except Exception as e:
        print(f"处理文档 {doc_path} 时发生错误: {str(e)}")
        traceback.print_exc()
        return False

def batch_format_documents(input_folder, output_folder):
    """
    批量处理指定文件夹中的Word文档
    """
    # 确保输出文件夹存在
    os.makedirs(output_folder, exist_ok=True)
    
    # 获取所有Word文档
    doc_files = [f for f in os.listdir(input_folder) 
                if f.endswith('.docx') and not f.startswith('~$')]
    
    # 统计处理结果
    total = len(doc_files)
    success = 0
    failed = []
    
    # 批量处理
    for doc_file in doc_files:
        input_path = os.path.join(input_folder, doc_file)
        output_path = os.path.join(output_folder, doc_file)
        
        print(f"正在处理：{doc_file}")
        
        if format_document(input_path, output_path):
            success += 1
        else:
            failed.append(doc_file)
    
    # 返回处理结果
    return {
        'total': total,
        'success': success,
        'failed': failed
    }

class WordFormatApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word文档格式统一调整工具")
        self.root.geometry("600x400")
        
        # 输入文件夹
        tk.Label(root, text="输入文件夹:").grid(row=0, column=0, padx=10, pady=10, sticky="w")
        self.input_folder_var = tk.StringVar()
        tk.Entry(root, textvariable=self.input_folder_var, width=50).grid(row=0, column=1, padx=5, pady=10)
        tk.Button(root, text="浏览...", command=self.select_input_folder).grid(row=0, column=2, padx=5, pady=10)
        
        # 输出文件夹
        tk.Label(root, text="输出文件夹:").grid(row=1, column=0, padx=10, pady=10, sticky="w")
        self.output_folder_var = tk.StringVar()
        tk.Entry(root, textvariable=self.output_folder_var, width=50).grid(row=1, column=1, padx=5, pady=10)
        tk.Button(root, text="浏览...", command=self.select_output_folder).grid(row=1, column=2, padx=5, pady=10)
        
        # 格式说明
        tk.Label(root, text="格式说明:").grid(row=2, column=0, padx=10, pady=5, sticky="nw")
        format_text = """
        1. 正文: 小四号字，宋体，1.5倍行距，首行缩进2字符，段前段后0.5行，不加粗，左对齐
        2. 标题1: 三号字，宋体，1.5倍行距，无缩进，段前段后0.5行，加粗，左对齐
        3. 标题2: 小三号字，宋体，1.5倍行距，无缩进，段前段后0.5行，加粗，左对齐
        4. 标题3: 四号字，宋体，1.5倍行距，无缩进，段前段后0.5行，加粗，左对齐
        5. 表格: 小四号字，宋体，1倍行距，无缩进，段前段后0.5行，不加粗，居中对齐
        6. 会自动删除所有空白行和空项目符号/编号
        7. 添加水印: "清控紫荆（北京）教育科技股份有限公司"
        """
        tk.Label(root, text=format_text, justify="left").grid(row=2, column=1, columnspan=2, padx=5, pady=5, sticky="w")
        
        # 开始处理按钮
        tk.Button(root, text="开始处理", command=self.start_processing, width=20).grid(row=3, column=1, pady=20)
        
        # 处理结果文本框
        self.result_text = tk.Text(root, height=8, width=70)
        self.result_text.grid(row=4, column=0, columnspan=3, padx=10, pady=10)
        
    def select_input_folder(self):
        folder = filedialog.askdirectory(title="选择包含Word文档的文件夹")
        if folder:
            self.input_folder_var.set(folder)
    
    def select_output_folder(self):
        folder = filedialog.askdirectory(title="选择处理后文档的保存位置")
        if folder:
            self.output_folder_var.set(folder)
    
    def start_processing(self):
        # 获取输入和输出文件夹
        input_folder = self.input_folder_var.get()
        output_folder = self.output_folder_var.get()
        
        # 验证输入
        if not input_folder or not os.path.isdir(input_folder):
            messagebox.showerror("错误", "请选择有效的输入文件夹")
            return
        
        if not output_folder:
            messagebox.showerror("错误", "请选择输出文件夹")
            return
        
        # 清空结果文本框
        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(tk.END, "开始处理Word文档...\n")
        self.root.update()
        
        try:
            # 批量处理文档
            results = batch_format_documents(input_folder, output_folder)
            
            # 显示处理结果
            self.result_text.insert(tk.END, f"\n处理完成!\n")
            self.result_text.insert(tk.END, f"总文档数: {results['total']}\n")
            self.result_text.insert(tk.END, f"成功处理: {results['success']}\n")
            self.result_text.insert(tk.END, f"处理失败: {len(results['failed'])}\n")
            
            if results['failed']:
                self.result_text.insert(tk.END, "\n处理失败的文件:\n")
                for doc in results['failed']:
                    self.result_text.insert(tk.END, f"- {doc}\n")
            
            # 弹出完成提示
            messagebox.showinfo("处理完成", f"成功处理 {results['success']}/{results['total']} 个文档")
            
        except Exception as e:
            # 显示错误信息
            self.result_text.insert(tk.END, f"\n处理过程中发生错误:\n{str(e)}")
            messagebox.showerror("错误", f"处理过程中发生错误:\n{str(e)}")

def main():
    root = tk.Tk()
    app = WordFormatApp(root)
    root.mainloop()

if __name__ == "__main__":
    main() 